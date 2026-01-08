import ollama
import docx
from docx import Document
from docx.shared import Pt, RGBColor
import fitz
import os
import logging
import argparse

# ----------------------------
#  Configuration
# ----------------------------
OUTPUT_FOLDER = "output"

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# ----------------------------
# 📚 Read and Extract Text
# ----------------------------
def read_resume(file_path):
    """
    Read and extract text from PDF, DOCX, or TXT files.
    """
    try:
        if file_path.endswith('.pdf'):
            return read_pdf(file_path)
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding="utf-8") as file:
                return file.read()
        else:
            logging.error(f"Unsupported file format: {file_path}")
            return None
    except Exception as e:
        logging.error(f"Error reading resume: {e}")
        return None


# ----------------------------
# 📖 Read PDF Files
# ----------------------------
def read_pdf(file_path):
    """
    Extract text from PDF using PyMuPDF.
    """
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text("text")
    return text.strip()


# ----------------------------
# 🤖 Generate Refined Resume with LLaMA
# ----------------------------
def generate_prompt(resume_text):
    """
    Generate an optimized, professional resume using LLaMA with ATS-friendly formatting.
    """
    prompt = f"""
    You are an expert resume writer with over 10 years of experience crafting ATS-optimized, result-driven, and professionally formatted resumes.

    Task:
    - Refine the following resume by removing unnecessary details and emphasizing quantifiable achievements.
    - Use strong action verbs, highlight relevant experience, and ensure industry-specific keywords are incorporated.
    - Format the resume with clean sections: Summary, Skills, Professional Experience, Education, Certifications, and optionally, Projects.
    - Maintain a clean, ATS-compatible structure using consistent bullet points and headings.
    - Use Markdown format and ensure the resume is aesthetically pleasing and professional.

    Resume Text:
    ```
    {resume_text}
    ```

    Provide the optimized resume in clean and structured Markdown format, ready to be converted into a professional DOCX.
    """
    try:
        response = ollama.chat(model="llama3.2:1b", messages=[{"role": "user", "content": prompt}])
        if 'message' in response and 'content' in response['message']:
            return response['message']['content']
        else:
            logging.error("Invalid response from LLaMA API.")
            return None
    except Exception as e:
        logging.error(f"Error generating refined resume: {e}")
        return None


# ----------------------------
# 🎨 Convert Markdown to DOCX with Advanced Formatting
# ----------------------------
def convert_markdown_to_docx(markdown_text, output_filename):
    """
    Convert Markdown text to a polished and beautifully formatted DOCX file.
    """
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in markdown_text.splitlines():
        line = line.strip()
        if line.startswith("# "):
            add_heading(doc, line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            add_heading(doc, line.replace("## ", ""), level=2)
        elif line.startswith("- "):
            add_bullet_point(doc, line.replace("- ", ""))
        elif line:
            add_paragraph(doc, line)

    # Set margins and improve layout
    set_doc_margins(doc)

    output_path = os.path.join(OUTPUT_FOLDER, output_filename)
    doc.save(output_path)
    logging.info(f"Refined resume successfully saved: {output_path}")
    return output_filename


# ----------------------------
# ✨ Helper Functions for DOCX Formatting
# ----------------------------
def add_heading(doc, text, level):
    """
    Add a heading with professional formatting.
    """
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.size = Pt(14) if level == 1 else Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for headings


def add_bullet_point(doc, text):
    """
    Add bullet points with consistent spacing and formatting.
    """
    paragraph = doc.add_paragraph(text, style="ListBullet")
    paragraph.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text):
    """
    Add normal paragraphs with appropriate spacing.
    """
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def set_doc_margins(doc):
    """
    Set consistent margins for the document.
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)


# ----------------------------
# 🚀 Main: Create Refined Resume
# ----------------------------
def create_refined_resume(file_path):
    """
    Main function to create a refined resume from an input file.
    """
    resume_text = read_resume(file_path)
    if not resume_text or not resume_text.strip():
        logging.warning("Resume file is empty or unreadable.")
        return None

    logging.info("Generating ATS-optimized resume...")
    refined_markdown = generate_prompt(resume_text)
    if not refined_markdown:
        logging.error("Failed to generate a refined resume.")
        return None

    docx_filename = file_path.split("/")[-1].replace(".pdf", ".docx").replace(".txt", ".docx").replace(".docx", "_refined.docx")
    return convert_markdown_to_docx(refined_markdown, docx_filename)


# ----------------------------
# 🎯 Add Command-Line Interface (CLI)
# ----------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Refine and optimize resumes for ATS.")
    parser.add_argument("file_path", help="Path to the resume file (PDF, DOCX, TXT)")
    args = parser.parse_args()

    result = create_refined_resume(args.file_path)
    if result:
        logging.info(f"Refined resume created successfully: {result}")
    else:
        logging.error("Failed to create a refined resume.")
