import ollama  # Ensure you have the Ollama Python library installed
import docx
import os
import fitz  # PyMuPDF for PDF reading
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ----------------------------
# 📝 Extract Text from Different File Formats
# ----------------------------
def read_resume(file_path):
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding="utf-8") as file:
            return file.read()
    elif file_path.endswith('.pdf'):
        return read_pdf(file_path)
    else:
        raise ValueError("Unsupported file format. Only .docx, .txt, and .pdf are supported.")


# ----------------------------
#  Extract Text from PDF Files
# ----------------------------
def read_pdf(file_path):
    try:
        doc = fitz.open(file_path)  # Open the PDF file
        text = ""
        for page_num in range(doc.page_count):  # Loop through all pages
            page = doc.load_page(page_num)  # Load each page
            text += page.get_text("text")  # Extract plain text from the page
        if not text.strip():
            raise ValueError(f"Warning: No text extracted from PDF file: {file_path}")
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


# ----------------------------
# 🎯 Extract Name and Location
# ----------------------------
def extract_name_and_location(resume_text):
    lines = resume_text.splitlines()
    name = "John Doe"  # Default fallback if name not found
    location = "Unknown Location"

    # Attempt to detect name and location from first 3 lines
    for line in lines[:3]:
        if len(line.split()) >= 2:  # Check if it's a plausible name
            name = line.strip()
        if "," in line or " " in line:
            location = line.strip()

    return name, location


# ----------------------------
# 🤖 Generate Refined Markdown with LLaMA
# ----------------------------
def generate_prompt(resume_text, name, location):
    prompt = f"""
    Given the following resume text, create a **refined, professional, and impactful resume** in **Markdown format**.
    Ensure that:
    - The **language is concise** and **highlights achievements**.
    - Sections are properly structured using **Markdown headings**, bullet points, and formatting.
    - The **output should be fully formatted Markdown**, including:
      - `# {name} - Resume` for the main header
      - `##` for section headers (e.g., "Experience", "Education", "Skills")
      - `-` for bullet points (e.g., key skills, job responsibilities).
    
    Resume Text:
    ```
    {resume_text}
    ```

    Please return **ONLY** the Markdown-formatted resume without any additional explanation.
    """

    print(f"Sending to LLaMA: \n{prompt}")  # Debugging: print the prompt being sent

    # Use Ollama API to process the resume and generate a refined version of it
    response = ollama.chat(model="llama3.2:1b", messages=[{"role": "user", "content": prompt}])
    
    # Debugging: print the raw response from LLaMA
    print(f"Response from LLaMA: \n{response}")

    return response['message']['content']


# ----------------------------
# 💾 Save Output as Markdown
# ----------------------------
def save_markdown(output_text, filename="refined_resume.md"):
    with open(filename, "w", encoding="utf-8") as md_file:
        md_file.write(output_text)
    print(f"✅ Markdown file saved as: {filename}")


# ----------------------------
# 🎨 Enhanced Markdown to DOCX Conversion (Beautiful + ATS Friendly)
# ----------------------------
def convert_markdown_to_docx(markdown_text, output_filename="refined_resume.docx"):
    doc = Document()

    # Define standard fonts and sizes for ATS
    font_name = "Calibri"  # Safe for ATS
    font_size_heading = 14
    font_size_body = 11

    # Split markdown text by lines
    lines = markdown_text.splitlines()

    for line in lines:
        line = line.strip()

        # Handle Main Heading (Candidate Name and Resume Title)
        if line.startswith("# "):
            heading_text = line.replace("# ", "").strip()
            heading = doc.add_heading(heading_text, level=1)
            run = heading.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size_heading)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center main heading

        # Handle Section Headings (Experience, Education, Skills)
        elif line.startswith("## "):
            section_text = line.replace("## ", "").strip()
            heading = doc.add_heading(section_text, level=2)
            run = heading.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size_heading)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Handle Subheadings or Job Titles (optional)
        elif line.startswith("### "):
            subheading_text = line.replace("### ", "").strip()
            para = doc.add_paragraph()
            run = para.add_run(subheading_text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_body)

        # Handle Bullet Points (Skills, Achievements, etc.)
        elif line.startswith("- "):
            bullet_text = line.replace("- ", "").strip()
            para = doc.add_paragraph(bullet_text, style="ListBullet")
            run = para.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size_body)

        # Handle Plain Text (For paragraphs or summaries)
        elif line:
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size_body)

    # Set consistent spacing for ATS compliance
    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(6)  # Space after paragraphs
        para.paragraph_format.line_spacing = 1.15  # Set 1.15 line spacing

    # Save the DOCX file
    doc.save(output_filename)
    print(f"✅ DOCX file saved successfully as: {output_filename}")


# ----------------------------
# 🚀 Main Function: Create Refined Resume and Generate DOCX
# ----------------------------
def create_refined_resume(file_path):
    try:
        # Check if file exists before proceeding
        if not os.path.isfile(file_path):
            return f"Error: File not found - {file_path}"

        # Extract text from the resume
        resume_text = read_resume(file_path)
        if not resume_text.strip():
            return "Error: No valid resume content found."

        # Extract name and location from the text
        name, location = extract_name_and_location(resume_text)
        print(f"Detected Name: {name}, Location: {location}")

        # Generate refined resume with user details
        refined_resume = generate_prompt(resume_text, name, location)

        # Save the refined resume as a Markdown file
        save_markdown(refined_resume)

        # Convert the refined Markdown resume to Enhanced DOCX
        docx_filename = "refined_resume.docx"
        convert_markdown_to_docx(refined_resume, docx_filename)

        return refined_resume
    except Exception as e:
        return str(e)


# ----------------------------
# 🧪 Example Usage
# ----------------------------
file_path = 'resume.pdf'  # Replace with actual file path (.docx, .txt, or .pdf)
refined_resume = create_refined_resume(file_path)
print(f"Refined Resume Output: \n{refined_resume}")
